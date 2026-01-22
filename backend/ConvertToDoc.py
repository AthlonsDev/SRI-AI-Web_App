from docx import Document

def convert_to_doc(text, output_filename):
    doc = Document()
    doc.add_heading('Transcription', 0)
    doc.add_paragraph(text)
    doc.save(output_filename)

    return output_filename